from PIL import Image, ImageDraw
from texify.output import replace_katex_invalid
from texify.model.processor import load_processor
from texify.model.model import load_model
from texify.inference import batch_inference
import pypdfium2
import hashlib
from streamlit_drawable_canvas import st_canvas
import streamlit as st
import pandas as pd
import io
import os
from docx import Document
from docx.shared import Inches
import requests
from pix2text import Pix2Text
import torch

os.environ["PYTORCH_ENABLE_MPS_FALLBACK"] = "1"

MAX_WIDTH = 800
MAX_HEIGHT = 1000

# Ensure the uploads directory exists
os.makedirs('uploads', exist_ok=True)

@st.cache_resource()
def load_model_cached():
    return load_model()

@st.cache_resource()
def load_processor_cached():
    return load_processor()

@st.cache_data()
def infer_image(pil_image, bbox, temperature):
    input_img = pil_image.crop(bbox)
    model_output = batch_inference([input_img], model, processor, temperature=temperature)
    return model_output[0]

def open_pdf(pdf_file):
    return pypdfium2.PdfDocument(io.BytesIO(pdf_file))

@st.cache_data()
def get_page_image(pdf_file_path, page_num, dpi=96):
    with open(pdf_file_path, 'rb') as f:
        doc = pypdfium2.PdfDocument(f)
        renderer = doc.render(pypdfium2.PdfBitmap.to_pil, page_indices=[page_num - 1], scale=dpi / 72)
        png = list(renderer)[0]
        png_image = png.convert("RGB")
    return png_image

@st.cache_data()
def get_uploaded_image(file_path):
    return Image.open(file_path).convert("RGB")

def resize_image(pil_image):
    if pil_image is None:
        return
    pil_image.thumbnail((MAX_WIDTH, MAX_HEIGHT), Image.Resampling.LANCZOS)

@st.cache_data()
def page_count(pdf_file_path):
    with open(pdf_file_path, 'rb') as f:
        doc = pypdfium2.PdfDocument(f)
        return len(doc)

def get_canvas_hash(pil_image):
    return hashlib.md5(pil_image.tobytes()).hexdigest()

@st.cache_data()
def get_image_size(pil_image):
    if pil_image is None:
        return MAX_HEIGHT, MAX_WIDTH
    height, width = pil_image.height, pil_image.width
    return height, width

st.set_page_config(layout="wide")

top_message = """### Math OCR

After the model loads, upload an image or a pdf, then draw a box around the equation or text you want to OCR by clicking and dragging. Texify will convert it to Markdown with LaTeX math on the right.

If you have already cropped your image, select "OCR image" in the sidebar instead.
"""

st.markdown(top_message)
col1, col2 = st.columns([.7, .3])

model = load_model_cached()
processor = load_processor_cached()

in_file = st.sidebar.file_uploader("PDF file or image:", type=["pdf", "png", "jpg", "jpeg", "gif", "webp"])
if in_file is None:
    st.stop()

# Save the uploaded file to the server
file_path = os.path.join('uploads', in_file.name)
with open(file_path, 'wb') as f:
    f.write(in_file.getvalue())

filetype = in_file.type
whole_image = False
if "pdf" in filetype:
    page_count = page_count(file_path)
    page_number = st.sidebar.number_input(f"Page number out of {page_count}:", min_value=1, value=1, max_value=page_count)

    pil_image = get_page_image(file_path, page_number)
else:
    pil_image = get_uploaded_image(file_path)
    whole_image = st.sidebar.button("OCR image")

# Resize to max bounds
resize_image(pil_image)

temperature = st.sidebar.slider("Generation temperature:", min_value=0.0, max_value=1.0, value=0.0, step=0.05)

canvas_hash = get_canvas_hash(pil_image) if pil_image else "canvas"

with col1:
    # Create a canvas component
    canvas_result = st_canvas(
        fill_color="rgba(255, 165, 0, 0.1)",
        stroke_width=1,
        stroke_color="#FFAA00",
        background_color="#FFF",
        background_image=pil_image,
        update_streamlit=True,
        height=get_image_size(pil_image)[0],
        width=get_image_size(pil_image)[1],
        drawing_mode="rect",
        point_display_radius=0,
        key=canvas_hash,
    )

if canvas_result.json_data is not None or whole_image:
    objects = pd.json_normalize(canvas_result.json_data["objects"])
    bbox_list = None
    if objects.shape[0] > 0:
        boxes = objects[objects["type"] == "rect"][["left", "top", "width", "height"]]
        boxes["right"] = boxes["left"] + boxes["width"]
        boxes["bottom"] = boxes["top"] + boxes["height"]
        bbox_list = boxes[["left", "top", "right", "bottom"]].values.tolist()
    if whole_image:
        bbox_list = [(0, 0, pil_image.width, pil_image.height)]

    if bbox_list:
        inferences = [infer_image(pil_image, bbox, temperature) for bbox in bbox_list]

        with col2:
            for idx, inference in enumerate(reversed(inferences)):
                st.markdown(f"### {len(inferences) - idx}")
                katex_markdown = replace_katex_invalid(inference)
                st.markdown(katex_markdown)
                st.code(inference)
                st.divider()

            if st.button("Export to DOCX"):
                doc = Document()
                doc.add_heading('OCR Results', 0)

                for idx, inference in enumerate(reversed(inferences)):
                    doc.add_heading(f'Inference {len(inferences) - idx}', level=1)
                    # Add the Markdown (reformulated text)
                    doc.add_paragraph(replace_katex_invalid(inference))

                doc_path = 'OCR_Results.docx'
                doc.save(doc_path)

                with open(doc_path, 'rb') as f:
                    st.download_button(
                        label="Download DOCX",
                        data=f,
                        file_name=doc_path,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

MAX_WIDTH = 800
MAX_HEIGHT = 1000

def convert_strokes_to_image(canvas_json_data, canvas_width, canvas_height):
    # Create a blank white image
    image = Image.new('RGB', (canvas_width, canvas_height), color='white')
    draw = ImageDraw.Draw(image)

    # Process each stroke from the JSON data
    for obj in canvas_json_data:
        if obj['type'] == 'rect':
            left = int(obj['left'])
            top = int(obj['top'])
            right = int(left + obj['width'])
            bottom = int(top + obj['height'])
            draw.rectangle([left, top, right, bottom], fill=None, outline='black')

    return image

drawing_mode = st.sidebar.selectbox(
    "Drawing tool:",
    ("freedraw", "line", "rect", "circle", "transform", "polygon", "point"),
)
stroke_width = st.sidebar.slider("Stroke width: ", 1, 25, 3)
if drawing_mode == "point":
    point_display_radius = st.sidebar.slider("Point display radius: ", 1, 25, 3)
stroke_color = st.sidebar.color_picker("Stroke color hex: ")
bg_color = st.sidebar.color_picker("Background color hex: ", "#eee")
bg_image = st.sidebar.file_uploader("Background image:", type=["png", "jpg"])

realtime_update = st.sidebar.checkbox("Update in realtime", True)

# Create a canvas component
canvas_result = st_canvas(
    fill_color="rgba(255, 165, 0, 0.3)",  # Fixed fill color with some opacity
    stroke_width=stroke_width,
    stroke_color=stroke_color,
    background_color=bg_color,
    background_image=Image.open(bg_image) if bg_image else None,
    update_streamlit=realtime_update,
    height=150,
    drawing_mode=drawing_mode,
    point_display_radius=point_display_radius if drawing_mode == "point" else 0,
    display_toolbar=st.sidebar.checkbox("Display toolbar", True),
    key="full_app",
)


def convert_handwritten_to_latex(img_fp):
    try:
        device = torch.device( 'cpu')

        p2t = Pix2Text.from_config(device=device)

        latex_formula = p2t.recognize_formula(img_fp)
        return latex_formula
    except Exception as e:
        print(f"Error converting handwritten to LaTeX: {e}")
        return None

# Display the canvas image
def array_to_image(array):
    return Image.fromarray(array.astype('uint8'))

# Assuming canvas_result.image_data is the drawn image data as a NumPy array
if canvas_result.image_data is not None:
    st.image(canvas_result.image_data)

    # Option to convert drawn image to LaTeX
    if st.button('Convert to LaTeX'):
        try:
            # Convert NumPy array to PIL Image
            image_pil = array_to_image(canvas_result.image_data)

            # Convert handwritten math expression to LaTeX
            latex_formula = convert_handwritten_to_latex(image_pil)

            # Display LaTeX output
            st.write(f"LaTeX Formula: {latex_formula}")
        except Exception as e:
            st.error(f"Error converting handwritten to LaTeX: {e}")

# Display the JSON data as a DataFrame

with col2:
    with st.expander("Usage tips"):
        tips = """
        ### Usage tips
        - Don't make your boxes too small or too large.
        - The model is sensitive to how you draw the box around the text you want to OCR. If you get bad results, try selecting a slightly different box, or splitting the box into multiple.
        - You can try changing the temperature value on the left if you don't get good results.  This controls how "creative" the model is.
        - Sometimes KaTeX (the script used to convert latex )won't be able to render an equation (red error text), but it will still be valid LaTeX.  You can copy the LaTeX and render it elsewhere.
        """
        st.markdown(tips)
